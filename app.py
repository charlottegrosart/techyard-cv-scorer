import anthropic
import streamlit as st
import PyPDF2
from docx import Document
import io

client = anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])

st.set_page_config(page_title="TechYard CV Scorer", page_icon="🎯")
st.markdown("""
    <style>
    .stButton > button {
        background-color: #dd762e;
        color: white;
        border: none;
        font-weight: 500;
    }
    .stButton > button:hover {
        background-color: #c4651f;
        color: white;
    }
    h1, h2, h3 {
        color: #dd762e !important;
    }
    </style>
""", unsafe_allow_html=True)

st.image("https://img.icons8.com/fluency/96/goal.png", width=60)
st.title("TechYard CV Scorer")
st.write("Powered by AI — score any candidate against a role in seconds.")

st.divider()

col1, col2 = st.columns(2)

with col1:
    st.subheader("Job description")
    job_file = st.file_uploader("Upload a job description", type=["pdf", "docx"])
    job_description = ""
    if job_file:
        if job_file.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(job_file)
            for page in reader.pages:
                job_description += page.extract_text()
        elif job_file.name.endswith(".docx"):
            doc = Document(io.BytesIO(job_file.read()))
            for para in doc.paragraphs:
                job_description += para.text + "\n"
        st.success("Job description uploaded successfully!")

with col2:
    st.subheader("Candidate CV")
    uploaded_file = st.file_uploader("Upload a CV", type=["pdf", "docx"])
    cv_text = ""
    if uploaded_file:
        if uploaded_file.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                cv_text += page.extract_text()
        elif uploaded_file.name.endswith(".docx"):
            doc = Document(io.BytesIO(uploaded_file.read()))
            for para in doc.paragraphs:
                cv_text += para.text + "\n"
        st.success("CV uploaded successfully!")

st.divider()

if st.button("Score this candidate", use_container_width=True):
    if job_description and cv_text:
        with st.spinner("Analysing candidate fit..."):
            prompt = "Score this CV against this job description from 1-10 and explain why. Structure your response like this:\n\n"
            prompt += "## Overall score: X/10\n\n"
            prompt += "## Summary\n"
            prompt += "A 2-3 sentence overview of the candidate fit.\n\n"
            prompt += "## Strengths and gaps\n"
            prompt += "| Category | Score | Status | Notes |\n"
            prompt += "|----------|-------|--------|-------|\n"
            prompt += "| Technical skills | X/10 | ✅ or ⚠️ or ❌ | Brief note |\n"
            prompt += "| Experience | X/10 | ✅ or ⚠️ or ❌ | Brief note |\n"
            prompt += "| Leadership | X/10 | ✅ or ⚠️ or ❌ | Brief note |\n"
            prompt += "| Location fit | X/10 | ✅ or ⚠️ or ❌ | Brief note |\n\n"
            prompt += "## Detailed breakdown\n"
            prompt += "A thorough analysis of strengths, gaps, and specific skill alignment.\n\n"
            prompt += "## Recommendation\n"
            prompt += "A clear hire / interview / pass recommendation with reasons.\n\n"
            prompt += "Job:\n" + job_description + "\n\nCV:\n" + cv_text

            message = client.messages.create(
                model="claude-opus-4-5",
                max_tokens=2048,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            st.divider()
            st.subheader("Assessment result")
            st.markdown(message.content[0].text)
    else:
        st.warning("Please upload both a job description and a CV before scoring.")